"""
IRIS — Générateur DOCX ATS-Optimized
Export CV en format Word éditable avec style IDENTIQUE au PDF (pixel-perfect)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.logger import get_logger
from backend.pdf_cv import THEMES, _ClassicDark

log = get_logger("docx_cv")


def color_to_rgb(color_obj) -> tuple:
    """Convertit un objet Color ReportLab en tuple RGB (0-255)."""
    if hasattr(color_obj, 'red'):
        return (int(color_obj.red * 255), int(color_obj.green * 255), int(color_obj.blue * 255))
    return (0, 0, 0)


def hex_to_rgb(hex_color: str) -> tuple:
    """Convertit #RRGGBB en tuple (R, G, B)."""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 6:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return (0, 0, 0)


def get_theme_colors(theme_name: str, custom_style: dict) -> dict:
    """Extrait les couleurs EXACTES du thème PDF pour application au DOCX."""
    theme_cls = THEMES.get(theme_name, _ClassicDark)
    
    # Si custom_style existe ET a des valeurs, utiliser ces couleurs (priorité)
    if custom_style and any(custom_style.values()):
        return {
            'name': hex_to_rgb(custom_style.get('name_color')) if custom_style.get('name_color') else color_to_rgb(getattr(theme_cls, 'WHITE', theme_cls.TEXT_LIGHT)),
            'title': hex_to_rgb(custom_style.get('subheading_color')) if custom_style.get('subheading_color') else color_to_rgb(getattr(theme_cls, 'TEXT_MED', theme_cls.TEXT_LIGHT)),
            'heading': hex_to_rgb(custom_style.get('heading_color')) if custom_style.get('heading_color') else color_to_rgb(theme_cls.ACCENT),
            'subheading': hex_to_rgb(custom_style.get('subheading_color')) if custom_style.get('subheading_color') else color_to_rgb(getattr(theme_cls, 'TEXT_MED', theme_cls.TEXT_LIGHT)),
            'text': hex_to_rgb(custom_style.get('text_color')) if custom_style.get('text_color') else color_to_rgb(theme_cls.TEXT_LIGHT),
            'accent': hex_to_rgb(custom_style.get('accent_color')) if custom_style.get('accent_color') else color_to_rgb(theme_cls.ACCENT),
        }
    
    # Sinon, extraire les couleurs du thème PDF
    colors = {
        'name': color_to_rgb(getattr(theme_cls, 'WHITE', theme_cls.TEXT_LIGHT)),
        'title': color_to_rgb(getattr(theme_cls, 'TEXT_MED', theme_cls.TEXT_LIGHT)),
        'heading': color_to_rgb(theme_cls.ACCENT),
        'subheading': color_to_rgb(getattr(theme_cls, 'TEXT_MED', theme_cls.TEXT_LIGHT)),
        'text': color_to_rgb(theme_cls.TEXT_LIGHT),
        'accent': color_to_rgb(theme_cls.ACCENT),
    }
    
    return colors


def generate_cv_docx(cv_data: dict) -> bytes:
    """
    Génère un CV au format DOCX avec le MÊME style visuel que le PDF.
    Applique les couleurs EXACTES du thème choisi (pixel-perfect).
    """
    theme = cv_data.get('theme', 'Classic Dark')
    custom_style = cv_data.get('custom_style', {})
    
    log.info(f"generate_cv_docx | name={cv_data.get('name')} | theme={theme}")
    
    # Extraire les couleurs EXACTES du thème PDF
    colors = get_theme_colors(theme, custom_style)
    
    doc = Document()
    
    # Marges identiques au PDF
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # 1. HEADER - Name & Title (COULEURS EXACTES DU PDF)
    name = cv_data.get("name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(name).upper())
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*colors['name'])
    
    title = cv_data.get("title", "")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(title))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*colors['title'])
    
    # 2. CONTACT INFO (COULEUR TEXTE DU PDF)
    contact_parts = []
    for key in ["email", "phone", "location"]:
        val = cv_data.get(key)
        if val:
            contact_parts.append(str(val))
    
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*colors['text'])
    
    # Links
    link_parts = []
    for key in ["linkedin", "github", "portfolio"]:
        val = cv_data.get(key)
        if val:
            link_parts.append(str(val))
    
    if link_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(link_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*colors['text'])
    
    doc.add_paragraph()
    
    # 3. SUMMARY
    summary = cv_data.get("summary")
    if summary:
        _add_section_title(doc, "PROFESSIONAL SUMMARY" if cv_data.get("lang") == "en" else "RÉSUMÉ PROFESSIONNEL", colors['heading'])
        p = doc.add_paragraph(str(summary))
        _format_body_text(p, colors['text'])
        doc.add_paragraph()
    
    # 4. EXPERIENCE (COULEURS EXACTES DU PDF)
    experiences = cv_data.get("experiences", [])
    if experiences:
        _add_section_title(doc, "PROFESSIONAL EXPERIENCE" if cv_data.get("lang") == "en" else "EXPÉRIENCE PROFESSIONNELLE", colors['heading'])
        
        for exp in experiences:
            role = exp.get("role", "")
            if role:
                p = doc.add_paragraph()
                run = p.add_run(str(role))
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*colors['subheading'])
            
            company_line = exp.get("company", "")
            location = exp.get("location", "")
            period = exp.get("period", "")
            
            if location:
                company_line += f" | {location}"
            if period:
                company_line += f" | {period}"
            
            if company_line:
                p = doc.add_paragraph(str(company_line))
                run = p.runs[0]
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(*colors['text'])
            
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                if bullet:
                    p = doc.add_paragraph(str(bullet), style='List Bullet')
                    _format_body_text(p, colors['text'])
            
            doc.add_paragraph()
    
    # 5. EDUCATION
    education = cv_data.get("education", [])
    if education:
        _add_section_title(doc, "EDUCATION" if cv_data.get("lang") == "en" else "FORMATION", colors['heading'])
        
        for edu in education:
            degree = edu.get("degree", "")
            if degree:
                p = doc.add_paragraph()
                run = p.add_run(str(degree))
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*colors['subheading'])
            
            school_line = edu.get("school", "")
            year = edu.get("year", "")
            if year:
                school_line += f" | {year}"
            
            if school_line:
                p = doc.add_paragraph(str(school_line))
                run = p.runs[0]
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(*colors['text'])
            
            detail = edu.get("detail")
            if detail:
                p = doc.add_paragraph(str(detail))
                _format_body_text(p, colors['text'])
            
            doc.add_paragraph()
    
    # 6. SKILLS
    skills = cv_data.get("skills", {})
    categories = skills.get("categories", [])
    if categories:
        _add_section_title(doc, "SKILLS & EXPERTISE" if cv_data.get("lang") == "en" else "COMPÉTENCES", colors['heading'])
        
        for cat in categories:
            cat_name = cat.get("name", "")
            items = cat.get("items", [])
            
            if cat_name and items:
                p = doc.add_paragraph()
                run = p.add_run(f"{cat_name}: ")
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(*colors['subheading'])
                
                run = p.add_run(", ".join(str(i) for i in items))
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(*colors['text'])
        
        doc.add_paragraph()
    
    # 7. LANGUAGES
    languages = cv_data.get("languages", [])
    if languages:
        _add_section_title(doc, "LANGUAGES" if cv_data.get("lang") == "en" else "LANGUES", colors['heading'])
        
        for lang in languages:
            lang_name = lang.get('lang', '')
            level = lang.get('level', '')
            if lang_name:
                lang_text = f"{lang_name} - {level}" if level else lang_name
                p = doc.add_paragraph(str(lang_text), style='List Bullet')
                _format_body_text(p, colors['text'])
        
        doc.add_paragraph()
    
    # 8. CERTIFICATIONS
    certifications = cv_data.get("certifications", [])
    if certifications:
        _add_section_title(doc, "CERTIFICATIONS", colors['heading'])
        
        for cert in certifications:
            if cert:
                p = doc.add_paragraph(str(cert), style='List Bullet')
                _format_body_text(p, colors['text'])
        
        doc.add_paragraph()
    
    # 9. INTERESTS
    interests = cv_data.get("interests", [])
    if interests:
        _add_section_title(doc, "INTERESTS" if cv_data.get("lang") == "en" else "CENTRES D'INTÉRÊT", colors['heading'])
        
        p = doc.add_paragraph(", ".join(str(i) for i in interests))
        _format_body_text(p, colors['text'])
    
    # Save to bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    result = buffer.getvalue()
    log.info(f"generate_cv_docx | done | size={len(result)} bytes")
    return result


def _add_section_title(doc, title: str, color: tuple):
    """Ajoute un titre de section avec la couleur EXACTE du thème PDF."""
    p = doc.add_paragraph()
    run = p.add_run(str(title))
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    run.font.underline = True


def _format_body_text(paragraph, color: tuple):
    """Formate le texte du corps avec la couleur EXACTE du thème PDF."""
    for run in paragraph.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*color)
